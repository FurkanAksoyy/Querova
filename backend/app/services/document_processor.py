import os
import uuid
import re
from typing import List, Dict, Any, Tuple, Optional
from pathlib import Path
import fitz  # PyMuPDF
from docx import Document
import logging

from app.config import settings
from app.models.document import DocumentChunk, DocumentMetadata

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Handles document parsing, chunking, and text extraction.

    Supports: PDF, DOCX, TXT
    Features:
    - Multi-format parsing
    - Semantic chunking with overlap
    - Metadata extraction (page numbers, positions)
    - Error handling and logging
    """

    def __init__(self):
        self.chunk_size = settings.chunk_size
        self.chunk_overlap = settings.chunk_overlap
        self.upload_dir = Path(settings.upload_directory)
        self.upload_dir.mkdir(exist_ok=True)
        logger.info(f"DocumentProcessor initialized: chunk_size={self.chunk_size}, overlap={self.chunk_overlap}")

    async def process_document(
            self,
            file_path: str,
            document_id: str
    ) -> Tuple[List[DocumentChunk], DocumentMetadata]:
        """
        Process a document: extract text, create chunks, and generate metadata.

        Args:
            file_path: Path to the uploaded document
            document_id: Unique identifier for the document

        Returns:
            Tuple of (list of document chunks, document metadata)

        Raises:
            ValueError: If file type is unsupported
            Exception: If parsing fails
        """
        file_path_obj = Path(file_path)
        file_extension = file_path_obj.suffix.lower()

        logger.info(f"Processing document: {document_id}, type: {file_extension}")

        try:
            # Extract text based on file type
            if file_extension == '.pdf':
                text, page_count = await self._extract_from_pdf(file_path)
            elif file_extension == '.docx':
                text, page_count = await self._extract_from_docx(file_path)
            elif file_extension == '.txt':
                text, page_count = await self._extract_from_txt(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")

            # Validate extracted text
            if not text or len(text.strip()) < 10:
                raise ValueError("Document appears to be empty or too short")

            # Create chunks
            chunks = self._create_chunks(text, document_id)

            if not chunks:
                raise ValueError("No chunks could be created from document")

            # Create metadata
            metadata = DocumentMetadata(
                filename=file_path_obj.name,
                file_type=file_extension,
                file_size=file_path_obj.stat().st_size,
                page_count=page_count,
                total_chunks=len(chunks)
            )

            logger.info(
                f"Document processed successfully: {document_id}, "
                f"{len(chunks)} chunks, {page_count} pages"
            )

            return chunks, metadata

        except Exception as e:
            logger.error(f"Failed to process document {document_id}: {str(e)}")
            raise

    async def _extract_from_pdf(self, file_path: str) -> Tuple[str, int]:
        """
        Extract text from PDF file.

        Args:
            file_path: Path to PDF file

        Returns:
            Tuple of (extracted text, page count)
        """
        try:
            doc = fitz.open(file_path)
            page_count = len(doc)  # Get count BEFORE closing
            text_parts = []

            for page_num in range(page_count):
                page = doc[page_num]
                text = page.get_text()

                # Clean up excessive whitespace
                text = re.sub(r'\n\s*\n', '\n\n', text)
                text = text.strip()

                if text:  # Only add non-empty pages
                    text_parts.append(f"[Page {page_num + 1}]\n{text}")

            doc.close()

            full_text = "\n\n".join(text_parts)
            logger.info(f"PDF extracted: {page_count} pages, {len(full_text)} characters")

            return full_text, page_count

        except Exception as e:
            logger.error(f"PDF extraction failed: {str(e)}")
            raise

    async def _extract_from_docx(self, file_path: str) -> Tuple[str, int]:
        """
        Extract text from DOCX file.

        Args:
            file_path: Path to DOCX file

        Returns:
            Tuple of (extracted text, estimated page count)
        """
        try:
            doc = Document(file_path)

            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        paragraphs.append(row_text)

            full_text = "\n\n".join(paragraphs)

            # Estimate page count (rough: 3000 chars per page)
            page_count = max(1, len(full_text) // 3000)

            logger.info(f"DOCX extracted: ~{page_count} pages, {len(full_text)} characters")

            return full_text, page_count

        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            raise

    async def _extract_from_txt(self, file_path: str) -> Tuple[str, int]:
        """
        Extract text from TXT file.

        Args:
            file_path: Path to TXT file

        Returns:
            Tuple of (extracted text, estimated page count)
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()

            # Estimate page count
            page_count = max(1, len(text) // 3000)

            logger.info(f"TXT extracted: ~{page_count} pages, {len(text)} characters")

            return text, page_count

        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    text = f.read()
                page_count = max(1, len(text) // 3000)
                return text, page_count
            except Exception as e:
                logger.error(f"TXT extraction failed with all encodings: {str(e)}")
                raise
        except Exception as e:
            logger.error(f"TXT extraction failed: {str(e)}")
            raise

    def _create_chunks(self, text: str, document_id: str) -> List[DocumentChunk]:
        """
        Split text into overlapping chunks using sliding window.

        Args:
            text: Full text to chunk
            document_id: Document identifier

        Returns:
            List of DocumentChunk objects

        Strategy:
        - Split by words (preserves semantic meaning better than chars)
        - Use sliding window with overlap
        - Extract page numbers from [Page X] markers
        - Store position metadata
        """
        chunks = []
        words = text.split()

        if len(words) == 0:
            logger.warning(f"No words found in document {document_id}")
            return chunks

        # Calculate chunk size in words (rough estimate: 1 word ≈ 5 chars)
        chunk_size_words = self.chunk_size // 5
        overlap_words = self.chunk_overlap // 5

        # Ensure reasonable values
        chunk_size_words = max(50, chunk_size_words)  # Min 50 words
        overlap_words = max(10, overlap_words)  # Min 10 words overlap

        start = 0
        chunk_index = 0

        while start < len(words):
            # Extract chunk
            end = min(start + chunk_size_words, len(words))
            chunk_words = words[start:end]
            chunk_text = " ".join(chunk_words)

            # Extract page number from [Page X] markers
            page_number = self._extract_page_number(chunk_text)

            # Create chunk
            chunk = DocumentChunk(
                chunk_id=f"{document_id}_chunk_{chunk_index}",
                document_id=document_id,
                text=chunk_text,
                page_number=page_number,
                chunk_index=chunk_index,
                metadata={
                    "word_count": len(chunk_words),
                    "char_count": len(chunk_text),
                    "start_word": start,
                    "end_word": end,
                    "is_complete": (end == len(words))
                }
            )
            chunks.append(chunk)

            # Move window forward
            start += chunk_size_words - overlap_words
            chunk_index += 1

            # Safety check: prevent infinite loop
            if chunk_index > 10000:
                logger.error(f"Chunk creation exceeded safety limit for {document_id}")
                break

        logger.info(
            f"Created {len(chunks)} chunks for {document_id}, "
            f"avg {sum(c.metadata['word_count'] for c in chunks) // len(chunks)} words/chunk"
        )

        return chunks

    def _extract_page_number(self, text: str) -> Optional[int]:
        """
        Extract page number from text containing [Page X] markers.

        Args:
            text: Text to search

        Returns:
            Page number if found, None otherwise
        """
        match = re.search(r'\[Page (\d+)\]', text)
        if match:
            return int(match.group(1))
        return None

    def save_uploaded_file(self, file_content: bytes, filename: str) -> str:
        """
        Save uploaded file to disk with unique identifier.

        Args:
            file_content: Raw file bytes
            filename: Original filename

        Returns:
            Path to saved file
        """
        # Generate unique filename
        file_id = str(uuid.uuid4())
        file_extension = Path(filename).suffix
        file_path = self.upload_dir / f"{file_id}{file_extension}"

        # Save file
        with open(file_path, 'wb') as f:
            f.write(file_content)

        logger.info(f"File saved: {file_path} ({len(file_content)} bytes)")

        return str(file_path)

    def cleanup_file(self, file_path: str) -> bool:
        """
        Delete a file from disk (optional cleanup).

        Args:
            file_path: Path to file

        Returns:
            True if deleted successfully
        """
        try:
            Path(file_path).unlink(missing_ok=True)
            logger.info(f"File deleted: {file_path}")
            return True
        except Exception as e:
            logger.error(f"Failed to delete file {file_path}: {str(e)}")
            return False


# Singleton instance
document_processor = DocumentProcessor()